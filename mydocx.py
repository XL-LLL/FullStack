import docx
import re

def fenge():
    print("-------------------->")


#world其实是压缩包，可以先改后缀为zip，然后解压就可以看到word的真=真实面目
#底层存储是xml
mydoc = docx.Document("icon/mydoc.docx")
#获取段落
p1 = mydoc.paragraphs[0]
print(p1.text)
print(p1.style.name)
#小段 就是一个段落里，根据风格 换行等等 被分为小段
p1r = p1.runs
print(p1r[4].text)
print(p1r[1].bold)

for runs in p1r:
    print(runs.text)


#获取所有段落
for p in mydoc.paragraphs:
    print(p.text)



#段落和图片合并读取 第一步 找到某段对于的图片的id

for p in mydoc.paragraphs:

    print(p._p.xml)
    if "Graphic" in p._p.xml:
        picnum = re.findall(' <a:blip r:embed="(\w+)">', p._p.xml)
        print("->",picnum)
    else:
        print("->",p.text)

#段落和图片合并读取 第一步 找到图片的对应关系
from docx.parts.image import ImagePart
img_dict = {}
for item in mydoc.part.rels.values():
    if type(item.target_part) == ImagePart:#对应关系里有许多关系 我们要图片的关系
        img_dict[item.rId] = item.target_part.partname
print(img_dict)

#段落和图片合并读取 第三步 结合
for p in mydoc.paragraphs:
    #print(p._p.xml)
    if "Graphic" in p._p.xml:
        [picnum] = re.findall(' <a:blip r:embed="(\w+)">', p._p.xml)
        #print("->", picnum)
        print("->",img_dict[picnum])

    else:
        print("->",p.text)

#段落和图片合并读取 第四步 提取图片
import shutil  #解压
mydocpath = "icon/mydoc.docx"
mydocpatzip = "icon/mydoc.zip"
mydocpatunzip = "icon/mydoc"
shutil.copy(mydocpath, mydocpatzip)#拷贝 相当于重命名 换了后缀
shutil.unpack_archive(mydocpatzip, mydocpatunzip)
#接下来 只需要根据地址就可以找到图片了



#段落和表格独立读取
for ta in mydoc.tables:
    for row in ta.rows:
        for cell in row.cells:
            print(cell.text)

fenge()
# 段落和表格合并读取
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from docx.table import Table
from docx.text.paragraph import  Paragraph

body = mydoc.element.body
for child in body.iterchildren():
    #print(child)
   if type(child) == CT_Tbl:

       table = Table(child,body)
       table_list = []
       for row in table.rows:
         row_list = []
         for cell in row.cells:
             row_list.append(cell.text)

         table_list.append(row_list)
       print(table_list)
   elif type(child) == CT_P:
       paragraph = Paragraph(child,body)
       print(paragraph.text)




#写入
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
doc_new = docx.Document()
p1 = doc_new.add_paragraph(text="sasa")

p1r1 = p1.add_run(text="mamas")
p1r1.add_break()
p1r2 = p1.add_run(text="zazaz")
p1r2.font.size = Pt(22)
p1r2.font.color.rgb = RGBColor(0,0,255)


p2 = doc_new.add_paragraph(text="sssss",style='List Bullet')
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc_new.add_paragraph(text="sasa")
p3.paragraph_format.left_indent = Pt(20)
p3.paragraph_format.line_spacing = Pt(20)

p4 = doc_new.add_table(3,3)
for row in p4.rows:
    for cell in row.cells:
        cell.add_paragraph(text="sasa")


#修改

p1 = doc_new.paragraphs[0]
p1r1 = p1.runs[0]
p1r1.text = "22222"
p1.insert_paragraph_before(text="11111111")
p1r1.add_text("2525252")
p1r1.add_picture("./static/a.png")

p2 = p1.insert_paragraph_before(text="000000000")
p2.clear()#清空
p2 = p1.insert_paragraph_before(text="000000000")
doc_new._body._body.remove(p2._p)#删除
p2 = p1.insert_paragraph_before(text="000000000")
p2._element.getparent().remove(p2._p)#删除

doc_new.save("./icon/mydocnew.docx")

